import os
import csv
import json
import logging
import subprocess
import shutil
import asyncio
from typing import Dict, Any, Callable, Optional, Tuple
from app.converters.base import BaseConverter
from app.services.pdf_tools_service import PdfToolsService

logger = logging.getLogger("document_converter")

class DocumentConverter(BaseConverter):
    async def convert(
        self,
        input_path: str,
        output_path: str,
        source_ext: str,
        target_ext: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> Tuple[bool, Optional[str]]:
        source_ext = source_ext.lower().lstrip(".")
        target_ext = target_ext.lower().lstrip(".")

        try:
            if progress_callback:
                progress_callback(10, f"Starting document conversion {source_ext.upper()} -> {target_ext.upper()}...")

            # ----------------------------------------------------
            # 1. DOCX / DOC Conversions
            # ----------------------------------------------------
            if source_ext in ["docx", "doc"]:
                if target_ext == "pdf":
                    # Attempt native LibreOffice conversion if installed in system
                    if shutil.which("libreoffice") or shutil.which("soffice"):
                        if progress_callback:
                            progress_callback(30, "Converting document to PDF via LibreOffice...")
                        if self._convert_native_libreoffice(input_path, output_path):
                            if progress_callback:
                                progress_callback(100, f"{source_ext.upper()} converted to PDF successfully.")
                            return True, None

                    if progress_callback:
                        progress_callback(30, "Extracting text, tables, and document elements...")

                    temp_img_dir = input_path + "_extracted_imgs"
                    os.makedirs(temp_img_dir, exist_ok=True)
                    footers_headers = {}
                    try:
                        elements, plain_paragraphs, footers_headers = self._extract_docx_elements(input_path, temp_img_dir)
                    except Exception as e:
                        logger.warning(f"Could not parse DOCX with python-docx, trying text fallback: {e}")
                        elements = []
                        plain_paragraphs = []
                        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                            for line in f:
                                if line.strip():
                                    elements.append({"type": "text", "content": line.strip()})
                                    plain_paragraphs.append(line.strip())

                    if progress_callback:
                        progress_callback(65, "Rendering layout, tables, and typography into PDF...")

                    success = self._render_elements_to_pdf(elements, output_path, footers_headers)
                    
                    # Clean up temp image folder
                    if os.path.exists(temp_img_dir):
                        shutil.rmtree(temp_img_dir, ignore_errors=True)

                    if success:
                        if progress_callback:
                            progress_callback(100, "DOCX converted to PDF successfully.")
                        return True, None
                    return False, "Failed to render PDF from Word document."

                elif target_ext in ["txt", "md"]:
                    # Clean up temp image folder
                    if os.path.exists(temp_img_dir):
                        shutil.rmtree(temp_img_dir, ignore_errors=True)
                    full_text = "\n\n".join(plain_paragraphs)
                    with open(output_path, "w", encoding="utf-8") as f_out:
                        f_out.write(full_text)
                    if progress_callback:
                        progress_callback(100, f"DOCX converted to {target_ext.upper()}.")
                    return True, None

            # ----------------------------------------------------
            # 2. TXT / MD Conversions
            # ----------------------------------------------------
            elif source_ext in ["txt", "md"]:
                with open(input_path, "r", encoding="utf-8", errors="replace") as f_in:
                    content = f_in.read()

                if target_ext == "pdf":
                    paragraphs = [line for line in content.splitlines() if line.strip()]
                    success = self._render_text_to_pdf(paragraphs, output_path)
                    if success:
                        if progress_callback:
                            progress_callback(100, f"{source_ext.upper()} converted to PDF.")
                        return True, None
                    return False, "Failed to render PDF from text file."

                elif target_ext in ["txt", "md"]:
                    with open(output_path, "w", encoding="utf-8") as f_out:
                        f_out.write(content)
                    if progress_callback:
                        progress_callback(100, "Text file converted.")
                    return True, None

                elif target_ext == "html":
                    html_lines = ["<!DOCTYPE html>", "<html>", "<head><meta charset='utf-8'><title>Document</title></head>", "<body style='font-family: sans-serif; padding: 2rem; max-width: 800px; margin: 0 auto;'>"]
                    for line in content.splitlines():
                        line_str = line.strip()
                        if not line_str:
                            continue
                        if line_str.startswith("# "):
                            html_lines.append(f"<h1>{line_str[2:]}</h1>")
                        elif line_str.startswith("## "):
                            html_lines.append(f"<h2>{line_str[3:]}</h2>")
                        elif line_str.startswith("### "):
                            html_lines.append(f"### {line_str[4:]}</h3>")
                        elif line_str.startswith("- ") or line_str.startswith("* "):
                            html_lines.append(f"<li>{line_str[2:]}</li>")
                        else:
                            html_lines.append(f"<p>{line_str}</p>")
                    html_lines.append("</body></html>")
                    with open(output_path, "w", encoding="utf-8") as f_out:
                        f_out.write("\n".join(html_lines))
                    if progress_callback:
                        progress_callback(100, "MD converted to HTML.")
                    return True, None

            # ----------------------------------------------------
            # 3. CSV Conversions
            # ----------------------------------------------------
            elif source_ext == "csv":
                rows = []
                with open(input_path, "r", encoding="utf-8", errors="replace") as f_in:
                    reader = csv.reader(f_in)
                    for row in reader:
                        rows.append(row)

                if target_ext == "xlsx":
                    try:
                        import openpyxl
                        wb = openpyxl.Workbook()
                        ws = wb.active
                        for r in rows:
                            ws.append(r)
                        wb.save(output_path)
                        if progress_callback:
                            progress_callback(100, "CSV converted to Excel XLSX.")
                        return True, None
                    except ImportError:
                        return False, "openpyxl library is not installed."

                elif target_ext == "json":
                    data = []
                    if len(rows) > 1:
                        headers = rows[0]
                        for r in rows[1:]:
                            obj = {headers[i]: r[i] if i < len(r) else "" for i in range(len(headers))}
                            data.append(obj)
                    with open(output_path, "w", encoding="utf-8") as f_out:
                        json.dump(data, f_out, indent=2)
                    if progress_callback:
                        progress_callback(100, "CSV converted to JSON.")
                    return True, None

                elif target_ext == "pdf":
                    formatted_rows = [" | ".join(r) for r in rows if any(r)]
                    success = self._render_text_to_pdf(formatted_rows, output_path)
                    if success:
                        if progress_callback:
                            progress_callback(100, "CSV converted to PDF.")
                        return True, None
                    return False, "Failed rendering CSV table to PDF."

            # ----------------------------------------------------
            # 4. JSON Conversions
            # ----------------------------------------------------
            elif source_ext == "json":
                with open(input_path, "r", encoding="utf-8", errors="replace") as f_in:
                    data = json.load(f_in)

                if isinstance(data, dict):
                    data = [data]

                if not isinstance(data, list) or len(data) == 0:
                    return False, "JSON file must contain a list of objects or a single key-value object."

                if target_ext == "csv":
                    headers = list(data[0].keys())
                    with open(output_path, "w", encoding="utf-8", newline="") as f_out:
                        writer = csv.DictWriter(f_out, fieldnames=headers)
                        writer.writeheader()
                        for item in data:
                            if isinstance(item, dict):
                                writer.writerow({k: str(v) for k, v in item.items() if k in headers})
                    if progress_callback:
                        progress_callback(100, "JSON converted to CSV.")
                    return True, None

            return False, f"Unsupported document conversion: {source_ext} -> {target_ext}"

        except Exception as e:
            logger.exception(f"Document conversion error: {e}")
            return False, f"Document conversion error: {str(e)}"

    async def _convert_native_word_com(self, input_path: str, output_path: str) -> bool:
        """Attempts native Windows Microsoft Word COM export for 1:1 pixel-perfect PDF conversion."""
        try:
            abs_input = os.path.abspath(input_path)
            abs_output = os.path.abspath(output_path)

            def export_com():
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                doc = word_app.Documents.Open(abs_input, ReadOnly=True)
                # 17 = wdFormatPDF
                doc.SaveAs2(abs_output, FileFormat=17)
                doc.Close()
                word_app.Quit()

            await asyncio.to_thread(export_com)
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                return True
            return False
        except Exception as e:
            logger.warning(f"Windows MS Word COM export attempt failed: {e}")
            return False

    def _extract_docx_elements(self, input_path: str, temp_dir: str) -> Tuple[list, list, Dict[str, str]]:
        """Extracts text paragraphs, tables, inline/floating images, and footers/headers in layout order."""
        import docx
        import zipfile
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        elements = []
        plain_paragraphs = []
        footers_headers = {"footer_left": "", "footer_right": "", "header": ""}
        extracted_media = set()
        img_counter = 0

        # Step 1: Open docx as a zip package and extract any raw images in word/media/
        media_map = {}
        try:
            with zipfile.ZipFile(input_path, 'r') as zf:
                for item in zf.namelist():
                    if item.startswith('word/media/') and not item.endswith('/'):
                        img_counter += 1
                        ext = item.rsplit('.', 1)[-1].lower() if '.' in item else 'png'
                        img_filename = f"extracted_img_{img_counter}.{ext}"
                        img_dest = os.path.join(temp_dir, img_filename)
                        with open(img_dest, 'wb') as f_out:
                            f_out.write(zf.read(item))
                        media_map[item] = img_dest
                        media_map[os.path.basename(item)] = img_dest
        except Exception as e:
            logger.warning(f"Failed extracting raw zip media from DOCX: {e}")

        try:
            doc = docx.Document(input_path)
        except Exception as e:
            logger.warning(f"python-docx could not open {input_path}: {e}")
            for img_path in media_map.values():
                elements.append({"type": "image", "path": img_path})
            return elements, plain_paragraphs, footers_headers

        # Extract section headers & footers
        for section in doc.sections:
            if section.footer and section.footer.paragraphs:
                f_texts = [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
                if f_texts:
                    full_f = " ".join(f_texts)
                    parts = [p.strip() for p in full_f.split("\t") if p.strip()]
                    if len(parts) >= 2:
                        footers_headers["footer_left"] = parts[0]
                        footers_headers["footer_right"] = parts[-1]
                    elif len(parts) == 1:
                        footers_headers["footer_left"] = parts[0]
            if section.header and section.header.paragraphs:
                h_texts = [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
                if h_texts:
                    footers_headers["header"] = " ".join(h_texts)

        def extract_images_from_element(elem):
            nonlocal img_counter
            rids = elem.xpath('.//*[local-name()="blip"]/@*[local-name()="embed"] | .//*[local-name()="imagedata"]/@*[local-name()="id" or local-name()="href"]')
            for rid in rids:
                try:
                    if rid in doc.part.related_parts:
                        img_part = doc.part.related_parts[rid]
                        img_bytes = img_part.image.blob
                        ext = img_part.image.ext or "png"
                        img_counter += 1
                        img_filename = f"extracted_img_{img_counter}.{ext}"
                        img_path = os.path.join(temp_dir, img_filename)
                        with open(img_path, "wb") as f_img:
                            f_img.write(img_bytes)
                        extracted_media.add(img_part.partname)
                        elements.append({"type": "image", "path": img_path})
                except Exception as e:
                    logger.warning(f"Failed extracting inline image rId {rid}: {e}")

        for child in doc.element.body:
            if child.tag.endswith('p'):
                p = Paragraph(child, doc)
                extract_images_from_element(child)
                text = p.text.strip()
                if text:
                    elements.append({"type": "text", "content": text})
                    plain_paragraphs.append(text)
            elif child.tag.endswith('tbl'):
                tbl = Table(child, doc)
                extract_images_from_element(child)
                table_matrix = []
                for row in tbl.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_matrix.append(row_cells)
                    row_txt = " | ".join(c for c in row_cells if c)
                    if row_txt:
                        plain_paragraphs.append(row_txt)
                if table_matrix:
                    elements.append({"type": "table", "matrix": table_matrix})

        # Append any remaining images from word/media that were not matched by XML IDs
        for key, img_dest in media_map.items():
            if not any(el.get("path") == img_dest for el in elements if el.get("type") == "image"):
                elements.append({"type": "image", "path": img_dest})

        return elements, plain_paragraphs, footers_headers

    def _render_elements_to_pdf(self, elements: list, output_path: str, footers_headers: Optional[dict] = None) -> bool:
        """Renders structured text blocks, tables, images, and headers/footers into clean A4 PDF pages."""
        try:
            import fitz
            from PIL import Image as PILImage

            pdf = fitz.open()
            width, height = 595.28, 841.89  # A4 standard points
            margin_left = 40.0
            margin_right = 40.0
            margin_top = 45.0
            margin_bottom = 55.0
            content_width = width - margin_left - margin_right

            footers = footers_headers or {}
            footer_left = footers.get("footer_left", "")
            footer_right = footers.get("footer_right", "")
            header_title = footers.get("header", "")

            pages = []

            def create_page():
                page = pdf.new_page(width=width, height=height)
                pages.append(page)

                if header_title:
                    h_w = fitz.get_text_length(header_title, fontsize=8.5)
                    page.insert_text((width - margin_right - h_w, 30), header_title, fontsize=8.5, color=(0.4, 0.4, 0.4))

                if footer_left or footer_right:
                    page.draw_line((margin_left, height - 42), (width - margin_right, height - 42), color=(0.85, 0.85, 0.85), width=0.5)
                    if footer_left:
                        page.insert_text((margin_left, height - 28), footer_left, fontsize=9, color=(0.2, 0.2, 0.2))
                    if footer_right:
                        fr_w = fitz.get_text_length(footer_right, fontsize=9)
                        page.insert_text((width - margin_right - fr_w, height - 28), footer_right, fontsize=9, color=(0.2, 0.2, 0.2))
                return page

            current_page = create_page()
            y = margin_top

            fontsize = 10.0
            line_height = 14.0

            for item in elements:
                if item["type"] == "text":
                    text_str = item["content"].strip()
                    if not text_str:
                        continue

                    chars_per_line = max(1, int(content_width / 5.2))
                    num_lines = max(1, (len(text_str) + chars_per_line - 1) // chars_per_line)
                    req_h = num_lines * line_height + 4

                    if y + req_h > height - margin_bottom:
                        current_page = create_page()
                        y = margin_top

                    rect = fitz.Rect(margin_left, y, width - margin_right, y + req_h + 10)
                    rc = current_page.insert_textbox(rect, text_str, fontsize=fontsize, fontname="helv", color=(0.1, 0.1, 0.1), align=0)

                    if rc < 0:
                        current_page = create_page()
                        y = margin_top
                        rect = fitz.Rect(margin_left, y, width - margin_right, height - margin_bottom)
                        current_page.insert_textbox(rect, text_str, fontsize=fontsize, fontname="helv", color=(0.1, 0.1, 0.1), align=0)
                        y += req_h + 6
                    else:
                        y += req_h + 6

                elif item["type"] == "table":
                    matrix = item.get("matrix", [])
                    if not matrix:
                        continue

                    num_rows = len(matrix)
                    num_cols = max(len(row) for row in matrix) if num_rows > 0 else 1
                    col_w = content_width / float(num_cols)
                    row_h = 24.0

                    tbl_height = num_rows * row_h
                    if y + tbl_height > height - margin_bottom:
                        current_page = create_page()
                        y = margin_top

                    for r_idx, row in enumerate(matrix):
                        for c_idx, cell_text in enumerate(row):
                            rx0 = margin_left + c_idx * col_w
                            ry0 = y + r_idx * row_h
                            rx1 = rx0 + col_w
                            ry1 = ry0 + row_h

                            cell_rect = fitz.Rect(rx0, ry0, rx1, ry1)
                            current_page.draw_rect(cell_rect, color=(0.2, 0.2, 0.2), width=0.75)

                            if cell_text:
                                text_rect = fitz.Rect(rx0 + 4, ry0 + 4, rx1 - 4, ry1 - 4)
                                current_page.insert_textbox(text_rect, cell_text, fontsize=9.5, fontname="helv", color=(0.1, 0.1, 0.1), align=0)

                    y += tbl_height + 10

                elif item["type"] == "image":
                    img_path = item["path"]
                    if not os.path.exists(img_path):
                        continue
                    try:
                        valid_img_path = img_path
                        img_w, img_h = 400.0, 300.0
                        try:
                            with PILImage.open(img_path) as pil_img:
                                img_w, img_h = float(pil_img.width), float(pil_img.height)
                                if pil_img.format not in ["JPEG", "PNG", "WEBP"]:
                                    png_path = img_path + "_converted.png"
                                    pil_img.convert("RGB").save(png_path, "PNG")
                                    valid_img_path = png_path
                        except Exception:
                            img_doc = fitz.open(img_path)
                            img_w, img_h = img_doc[0].rect.width, img_doc[0].rect.height
                            img_doc.close()

                        max_w = content_width
                        max_h = 380.0
                        scale = min(max_w / max(1.0, img_w), max_h / max(1.0, img_h), 1.0)
                        display_w = img_w * scale
                        display_h = img_h * scale

                        if y + display_h > height - margin_bottom:
                            current_page = create_page()
                            y = margin_top

                        img_rect = fitz.Rect(margin_left, y, margin_left + display_w, y + display_h)
                        current_page.insert_image(img_rect, filename=valid_img_path)
                        y += display_h + 10
                    except Exception as img_err:
                        logger.warning(f"Error rendering image in PDF: {img_err}")

            if len(pdf) == 0:
                create_page()

            pdf.save(output_path)
            pdf.close()
            PdfToolsService.remove_blank_pages(output_path, output_path)
            return True
        except Exception as e:
            logger.error(f"Error rendering elements layout to PDF: {e}")
            return False

    def _convert_native_libreoffice(self, input_path: str, output_path: str) -> bool:
        """Attempts 1:1 pixel-perfect native PDF conversion via LibreOffice CLI."""
        try:
            out_dir = os.path.dirname(os.path.abspath(output_path))
            abs_input = os.path.abspath(input_path)

            cmd = ["libreoffice", "--headless", "--convert-to", "pdf", abs_input, "--outdir", out_dir]
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=45)

            input_base = os.path.basename(abs_input).rsplit(".", 1)[0]
            generated_pdf = os.path.join(out_dir, f"{input_base}.pdf")

            if os.path.exists(generated_pdf):
                if os.path.abspath(generated_pdf) != os.path.abspath(output_path):
                    shutil.move(generated_pdf, output_path)
                PdfToolsService.remove_blank_pages(output_path, output_path)
                return True
            return False
        except Exception as e:
            logger.warning(f"LibreOffice conversion attempt failed: {e}")
            return False

    def _render_text_to_pdf(self, paragraphs: list, output_path: str) -> bool:
        """Renders structured paragraphs cleanly onto A4 PDF pages with auto line-wrapping and clean margins."""
        try:
            import fitz
            pdf = fitz.open()
            width, height = 595.28, 841.89  # A4 standard points
            margin_left = 40.0
            margin_right = 40.0
            margin_top = 45.0
            margin_bottom = 55.0
            content_width = width - margin_left - margin_right

            page = pdf.new_page(width=width, height=height)
            y = margin_top

            fontsize = 10.0
            line_height = 14.0

            for para in paragraphs:
                if not para or not str(para).strip():
                    y += line_height
                    continue

                text_str = str(para).strip()
                chars_per_line = max(1, int(content_width / 5.2))
                num_lines = max(1, (len(text_str) + chars_per_line - 1) // chars_per_line)
                req_h = num_lines * line_height + 4

                if y + req_h > height - margin_bottom:
                    page = pdf.new_page(width=width, height=height)
                    y = margin_top

                rect = fitz.Rect(margin_left, y, width - margin_right, y + req_h + 10)
                rc = page.insert_textbox(rect, text_str, fontsize=fontsize, fontname="helv", color=(0.1, 0.1, 0.1), align=0)
                if rc < 0:
                    page = pdf.new_page(width=width, height=height)
                    y = margin_top
                    rect = fitz.Rect(margin_left, y, width - margin_right, height - margin_bottom)
                    page.insert_textbox(rect, text_str, fontsize=fontsize, fontname="helv", color=(0.1, 0.1, 0.1), align=0)
                    y += req_h + 6
                else:
                    y += req_h + 6

            if len(pdf) == 0:
                pdf.new_page(width=width, height=height)

            pdf.save(output_path)
            pdf.close()
            PdfToolsService.remove_blank_pages(output_path, output_path)
            return True
        except Exception as e:
            logger.error(f"Error rendering PDF layout: {e}")
            return False
