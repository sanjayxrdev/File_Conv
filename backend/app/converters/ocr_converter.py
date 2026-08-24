"""
Docling OCR & Document Intelligence Engine
==========================================
Powered by Docling (https://github.com/docling-project/docling)

Provides AI-driven document layout recognition, table structure extraction,
formula parsing, reading order recovery, and OCR for scanned PDFs, images,
and digital office documents.
"""

import os
import json
import logging
import asyncio
from typing import Dict, Any, Callable, Optional, Tuple
from app.converters.base import BaseConverter

logger = logging.getLogger("ocr_converter")


class DoclingOCRConverter(BaseConverter):
    """
    OCR and Document Intelligence Converter using IBM Docling.
    
    Supported Source Formats:
      - PDF: Scanned & Digital (.pdf)
      - Images: (.png, .jpg, .jpeg, .webp, .tiff, .bmp)
      - Documents: (.docx, .pptx, .xlsx, .html)
      
    Supported Target Formats:
      - .md   : Clean structured Markdown with layout and tables
      - .txt  : Plain text document extraction
      - .json : Complete Docling Document AST with bounding boxes & structure
      - .html : Semantic HTML document with inline styling
      - .docx : Formatted Word document
    """

    def __init__(self):
        self._converter = None

    def _get_converter(self):
        """Lazy-loads and caches the Docling DocumentConverter instance."""
        if self._converter is None:
            try:
                from docling.document_converter import DocumentConverter
                self._converter = DocumentConverter()
                logger.info("Docling DocumentConverter initialized successfully.")
            except Exception as e:
                logger.error(f"Failed to initialize Docling DocumentConverter: {e}")
                raise
        return self._converter

    def _run_docling_sync(
        self,
        input_path: str,
        output_path: str,
        source_ext: str,
        target_ext: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> Tuple[bool, Optional[str]]:
        """Synchronous converter invocation intended to run in worker thread."""
        try:
            if progress_callback:
                progress_callback(15, "Initializing Docling OCR & Layout Parser...")

            converter = self._get_converter()

            if progress_callback:
                progress_callback(35, f"Analyzing {source_ext.upper()} layout and running OCR...")

            conv_result = converter.convert(input_path)

            if progress_callback:
                progress_callback(75, "Extracting text, tables, and structured document hierarchy...")

            doc = conv_result.document

            # ----------------------------------------------------
            # TARGET 1: Markdown (.md)
            # ----------------------------------------------------
            if target_ext == "md":
                md_content = doc.export_to_markdown()
                with open(output_path, "w", encoding="utf-8") as f_out:
                    f_out.write(md_content)

            # ----------------------------------------------------
            # TARGET 2: Plain Text (.txt)
            # ----------------------------------------------------
            elif target_ext == "txt":
                try:
                    text_content = doc.export_to_text()
                except Exception:
                    # Fallback to markdown text if export_to_text is not directly available
                    text_content = doc.export_to_markdown()
                with open(output_path, "w", encoding="utf-8") as f_out:
                    f_out.write(text_content)

            # ----------------------------------------------------
            # TARGET 3: Docling JSON AST (.json)
            # ----------------------------------------------------
            elif target_ext == "json":
                try:
                    json_data = doc.export_to_dict()
                except Exception:
                    json_data = {"document": str(doc)}
                with open(output_path, "w", encoding="utf-8") as f_out:
                    json.dump(json_data, f_out, indent=2, ensure_ascii=False)

            # ----------------------------------------------------
            # TARGET 4: Semantic HTML (.html)
            # ----------------------------------------------------
            elif target_ext == "html":
                try:
                    html_content = doc.export_to_html()
                except Exception:
                    md_text = doc.export_to_markdown()
                    html_content = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Docling OCR Result</title><style>body{{font-family:sans-serif;max-width:850px;margin:2rem auto;padding:1rem;line-height:1.6}}table{{border-collapse:collapse;width:100%}}th,td{{border:1px solid #ccc;padding:8px}}</style></head><body><pre>{md_text}</pre></body></html>"
                with open(output_path, "w", encoding="utf-8") as f_out:
                    f_out.write(html_content)

            # ----------------------------------------------------
            # TARGET 5: Word Document (.docx)
            # ----------------------------------------------------
            elif target_ext == "docx":
                import docx
                word_doc = docx.Document()
                md_text = doc.export_to_markdown()
                
                for line in md_text.splitlines():
                    stripped = line.strip()
                    if not stripped:
                        continue
                    if stripped.startswith("# "):
                        word_doc.add_heading(stripped[2:], level=1)
                    elif stripped.startswith("## "):
                        word_doc.add_heading(stripped[3:], level=2)
                    elif stripped.startswith("### "):
                        word_doc.add_heading(stripped[4:], level=3)
                    elif stripped.startswith("- ") or stripped.startswith("* "):
                        word_doc.add_paragraph(stripped[2:], style='List Bullet')
                    else:
                        word_doc.add_paragraph(stripped)
                
                word_doc.save(output_path)

            else:
                return False, f"Unsupported OCR target format: {target_ext}"

            if progress_callback:
                progress_callback(100, f"Docling OCR conversion to {target_ext.upper()} completed successfully.")

            return True, None

        except Exception as e:
            logger.exception(f"Docling conversion failed: {e}")
            return False, f"Docling OCR processing error: {str(e)}"

    async def convert(
        self,
        input_path: str,
        output_path: str,
        source_ext: str,
        target_ext: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> Tuple[bool, Optional[str]]:
        """
        Executes Docling OCR conversion asynchronously inside a thread pool.
        """
        source_ext = source_ext.lower().lstrip(".")
        target_ext = target_ext.lower().lstrip(".")

        if not os.path.exists(input_path):
            return False, f"Source file does not exist: {input_path}"

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self._run_docling_sync,
            input_path,
            output_path,
            source_ext,
            target_ext,
            options,
            progress_callback
        )
