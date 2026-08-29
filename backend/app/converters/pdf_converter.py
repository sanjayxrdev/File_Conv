import os
import io
import re
import logging
from typing import Dict, Any, Callable, Optional, Tuple
from app.converters.base import BaseConverter

logger = logging.getLogger("pdf_converter")

class PDFConverter(BaseConverter):
    async def convert(
        self,
        input_path: str,
        output_path: str,
        source_ext: str,
        target_ext: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> Tuple[bool, Optional[str]]:
        options = options or {}
        target_ext = target_ext.lower().lstrip(".")

        try:
            import fitz  # PyMuPDF
        except ImportError:
            return False, "PyMuPDF (fitz) is not installed in the Python environment."

        try:
            if progress_callback:
                progress_callback(10, "Opening PDF document...")

            try:
                doc = fitz.open(input_path)
            except Exception as pdf_err:
                # Handle text files saved with .pdf extension
                try:
                    with open(input_path, "r", encoding="utf-8", errors="ignore") as tf:
                        raw_text = tf.read()
                    if raw_text.strip():
                        if target_ext == "docx":
                            import docx
                            wdoc = docx.Document()
                            for para in raw_text.split("\n\n"):
                                if para.strip():
                                    wdoc.add_paragraph(para.strip())
                            wdoc.save(output_path)
                            if progress_callback:
                                progress_callback(100, "Text-based PDF converted to DOCX successfully.")
                            return True, None
                        elif target_ext in ["txt", "md"]:
                            with open(output_path, "w", encoding="utf-8") as out_f:
                                out_f.write(raw_text)
                            if progress_callback:
                                progress_callback(100, f"Text-based PDF saved to {target_ext.upper()}.")
                            return True, None
                except Exception:
                    pass
                return False, f"Invalid PDF file structure: {str(pdf_err)}"

            total_pages = len(doc)

            if total_pages == 0:
                doc.close()
                return False, "PDF document has 0 pages or is corrupted."

            # Check if PDF might be scanned (no extractable text)
            has_text = False
            for page in doc:
                if page.get_text().strip():
                    has_text = True
                    break

            if not has_text and target_ext in ["txt", "md", "docx"]:
                logger.warning("PDF appears to be scanned image with no embedded text.")

            # ----------------------------------------------------
            # 1. PDF -> TXT
            # ----------------------------------------------------
            if target_ext == "txt":
                text_content = []
                for idx, page in enumerate(doc):
                    text_content.append(f"--- Page {idx + 1} ---\n" + page.get_text())
                    if progress_callback:
                        pct = int(10 + (idx + 1) / total_pages * 80)
                        progress_callback(pct, f"Extracting page {idx + 1}/{total_pages}")

                full_text = "\n\n".join(text_content)
                if not has_text:
                    full_text = "[NOTE: Document appears scanned. OCR may be required for full text extraction.]\n\n" + full_text

                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(full_text)

                doc.close()
                if progress_callback:
                    progress_callback(100, "PDF to TXT conversion completed.")
                return True, None

            # ----------------------------------------------------
            # 2. PDF -> Markdown
            # ----------------------------------------------------
            elif target_ext == "md":
                md_content = []
                for idx, page in enumerate(doc):
                    lines = page.get_text("text").splitlines()
                    md_lines = [f"## Page {idx + 1}\n"]
                    for line in lines:
                        if line.isupper() and len(line) < 60:
                            md_lines.append(f"### {line}")
                        else:
                            md_lines.append(line)
                    md_content.append("\n".join(md_lines))
                    if progress_callback:
                        pct = int(10 + (idx + 1) / total_pages * 80)
                        progress_callback(pct, f"Processing page {idx + 1}/{total_pages}")

                full_md = "\n\n---\n\n".join(md_content)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(full_md)

                doc.close()
                if progress_callback:
                    progress_callback(100, "PDF to Markdown conversion completed.")
                return True, None

            # ----------------------------------------------------
            # 3. PDF -> Image (PNG/JPG)
            # ----------------------------------------------------
            elif target_ext in ["png", "jpg", "jpeg"]:
                dpi = int(options.get("dpi", 150))
                zoom = dpi / 72.0
                mat = fitz.Matrix(zoom, zoom)

                if total_pages == 1:
                    page = doc.load_page(0)
                    pix = page.get_pixmap(matrix=mat)
                    pix.save(output_path)
                    doc.close()
                    if progress_callback:
                        progress_callback(100, "PDF page rendered as image.")
                    return True, None
                else:
                    import zipfile
                    import tempfile

                    with tempfile.TemporaryDirectory() as tmpdir:
                        img_paths = []
                        for idx, page in enumerate(doc):
                            pix = page.get_pixmap(matrix=mat)
                            img_filename = f"page_{idx + 1:02d}.{target_ext}"
                            img_path = os.path.join(tmpdir, img_filename)
                            pix.save(img_path)
                            img_paths.append((img_path, img_filename))
                            if progress_callback:
                                pct = int(10 + (idx + 1) / total_pages * 80)
                                progress_callback(pct, f"Rendering page {idx + 1}/{total_pages} as image...")

                        # Package all page images into a ZIP archive
                        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zip_file:
                            for img_path, img_filename in img_paths:
                                zip_file.write(img_path, arcname=img_filename)

                    doc.close()
                    if progress_callback:
                        progress_callback(100, f"Rendered all {total_pages} pages to ZIP archive.")
                    return True, None

            # ----------------------------------------------------
            # 4. PDF -> DOCX / DOC (Universal High-Fidelity Conversion)
            # ----------------------------------------------------
            elif target_ext in ["docx", "doc"]:
                if progress_callback:
                    progress_callback(20, "Analyzing PDF page geometry, headers, footers & layout...")

                # If PDF has no extractable text (e.g. scanned pages), extract high-resolution page layout
                if not has_text:
                    if progress_callback:
                        progress_callback(40, "Scanned document detected. Extracting high-res page image layout...")
                    return self._fallback_image_render_docx(input_path, output_path, progress_callback)

                # Execute pure pdf2docx conversion directly for 100% native image and layout preservation
                try:
                    from pdf2docx import Converter
                    if progress_callback:
                        progress_callback(50, "Extracting document structure, tables, fonts & images...")

                    cv = Converter(input_path)
                    cv.convert(output_path)
                    cv.close()

                    if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                        # Ensure all embedded images are visible (fix behindDoc="1" hiding images behind text/canvas)
                        try:
                            import docx
                            word_doc = docx.Document(output_path)
                            anchors = word_doc._element.xpath('.//*[local-name()="anchor"]')
                            modified = False
                            for a in anchors:
                                if a.get('behindDoc') == '1':
                                    a.set('behindDoc', '0')
                                    a.set('relativeHeight', '251658240')
                                    modified = True
                            if modified:
                                word_doc.save(output_path)
                        except Exception as anchor_err:
                            logger.warning(f"Failed to normalize anchor visibility: {anchor_err}")

                        doc.close()
                        if progress_callback:
                            progress_callback(100, "PDF to Word conversion completed successfully.")
                        return True, None
                except Exception as p2d_err:
                    logger.warning(f"pdf2docx engine conversion failed: {p2d_err}, falling back to custom extractor.")

                # Fallback to high-fidelity extractor if pdf2docx encountered an issue
                doc.close()
                return self._fallback_image_render_docx(input_path, output_path, progress_callback)

            else:
                doc.close()
                return False, f"Unsupported target format for PDF: {target_ext}"

        except Exception as e:
            logger.exception(f"PDF conversion error: {e}")
            return False, f"PDF conversion error: {str(e)}"

    def _fallback_image_render_docx(
        self,
        input_path: str,
        output_path: str,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> Tuple[bool, Optional[str]]:
        """High-fidelity fallback extractor using PyMuPDF and python-docx for text and images."""
        try:
            import fitz
            import docx
            from docx.shared import Inches

            doc = fitz.open(input_path)
            word_doc = docx.Document()
            processed_xrefs = set()

            for section in word_doc.sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            total_pages = len(doc)
            for idx, page in enumerate(doc):
                if idx > 0:
                    word_doc.add_page_break()

                # Extract text blocks
                blocks = page.get_text("blocks")
                for b in blocks:
                    block_text = str(b[4]).strip()
                    if block_text:
                        word_doc.add_paragraph(block_text)

                # Extract embedded images on page
                images = page.get_images(full=True)
                for img_info in images:
                    try:
                        xref = img_info[0]
                        if xref in processed_xrefs:
                            continue
                        processed_xrefs.add(xref)

                        base_img = doc.extract_image(xref)
                        if base_img and "image" in base_img:
                            w = base_img.get("width", 100)
                            h = base_img.get("height", 100)
                            if w >= 20 and h >= 20:
                                img_bytes = base_img["image"]
                                img_stream = io.BytesIO(img_bytes)
                                p = word_doc.add_paragraph()
                                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                target_width = min(6.5, max(1.5, w / 96.0))
                                p.add_run().add_picture(img_stream, width=Inches(target_width))
                    except Exception as img_err:
                        logger.warning(f"Failed to extract image xref {img_info[0]}: {img_err}")

                # If no text blocks and no embedded images on page (scanned raster page), render page pixmap
                if not blocks and not images:
                    pix = page.get_pixmap(dpi=180)
                    img_bytes = pix.tobytes("png")
                    img_stream = io.BytesIO(img_bytes)
                    p = word_doc.add_paragraph()
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_stream, width=Inches(6.8))

                if progress_callback:
                    pct = int(50 + (idx + 1) / total_pages * 50)
                    progress_callback(pct, f"Extracting page {idx + 1}/{total_pages}...")

            doc.close()
            word_doc.save(output_path)
            if progress_callback:
                progress_callback(100, "PDF to DOCX conversion completed.")
            return True, None
        except Exception as e:
            return False, f"Fallback PDF rendering failed: {str(e)}"
