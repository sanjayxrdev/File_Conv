import os
import logging
from typing import Dict, Any, Callable, Optional, Tuple
from app.converters.base import BaseConverter
from app.services.pdf_tools_service import PdfToolsService

logger = logging.getLogger("html_converter")

class HTMLConverter(BaseConverter):
    async def convert(
        self,
        input_path: str,
        output_path: str,
        source_ext: str,
        target_ext: str,
        options: Optional[Dict[str, Any]] = None,
        progress_callback: Optional[Callable[[int, str], None]] = None
    ) -> Tuple[bool, Optional[str]]:
        target_ext = target_ext.lower().lstrip(".")

        try:
            if progress_callback:
                progress_callback(10, f"Reading HTML file...")

            with open(input_path, "r", encoding="utf-8", errors="replace") as f:
                html_content = f.read()

            # HTML -> Markdown (.md)
            if target_ext == "md":
                try:
                    import html2text
                    h = html2text.HTML2Text()
                    h.ignore_links = False
                    md_text = h.handle(html_content)
                    with open(output_path, "w", encoding="utf-8") as out:
                        out.write(md_text)
                    if progress_callback:
                        progress_callback(100, "HTML converted to Markdown.")
                    return True, None
                except Exception as e:
                    return False, f"HTML to MD error: {str(e)}"

            # HTML -> TXT
            elif target_ext == "txt":
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(html_content, "html.parser")
                    text = soup.get_text(separator="\n\n")
                    with open(output_path, "w", encoding="utf-8") as out:
                        out.write(text)
                    if progress_callback:
                        progress_callback(100, "HTML converted to Plain Text.")
                    return True, None
                except Exception as e:
                    return False, f"HTML to TXT error: {str(e)}"

            # HTML -> PDF
            elif target_ext == "pdf":
                try:
                    import fitz
                    if progress_callback:
                        progress_callback(40, "Rendering HTML to PDF layout with high-fidelity styles...")

                    css_style = """
                    <style>
                    @page { size: A4; margin: 15mm; }
                    html, body {
                        width: 100%;
                        font-family: Arial, Helvetica, sans-serif;
                        font-size: 10.5pt;
                        line-height: 1.45;
                        color: #111111;
                        margin: 0;
                        padding: 0;
                        box-sizing: border-box;
                    }
                    table {
                        width: 100% !important;
                        border-collapse: collapse !important;
                        margin-top: 8px;
                        margin-bottom: 12px;
                    }
                    th, td {
                        border: 1px solid #333333 !important;
                        padding: 6px 8px !important;
                        vertical-align: top;
                    }
                    pre, code {
                        font-family: "Courier New", Courier, monospace;
                        font-size: 9.5pt;
                        white-space: pre-wrap !important;
                        word-wrap: break-word !important;
                        word-break: break-all !important;
                        background: #f8f9fa;
                        border: 1px solid #e9ecef;
                        padding: 6px;
                        border-radius: 4px;
                    }
                    img {
                        max-width: 100% !important;
                        height: auto !important;
                    }
                    </style>
                    """

                    if "<head>" in html_content:
                        styled_html = html_content.replace("<head>", f"<head>{css_style}")
                    else:
                        styled_html = f"<html><head>{css_style}</head><body>{html_content}</body></html>"

                    story = fitz.Story(styled_html)
                    writer = fitz.DocumentWriter(output_path)
                    more = True
                    # A4 dimensions 595.28 x 841.89 pt; 36pt (0.5 in) margins
                    while more:
                        device = writer.begin_page(fitz.paper_rect("a4"))
                        more, _ = story.place(fitz.Rect(36, 36, 559.28, 805.89))
                        story.draw(device)
                        writer.end_page()
                    writer.close()

                    PdfToolsService.remove_blank_pages(output_path, output_path)

                    if progress_callback:
                        progress_callback(100, "HTML to PDF rendering complete.")
                    return True, None
                except Exception as e:
                    logger.warning(f"PyMuPDF Story HTML to PDF failed, fallback using text layout: {e}")
                    import fitz
                    doc = fitz.open()
                    page = doc.new_page(width=595.28, height=841.89)
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(html_content, "html.parser")
                    page.insert_textbox(fitz.Rect(36, 36, 559.28, 805.89), soup.get_text(), fontsize=10.5)
                    doc.save(output_path)
                    doc.close()
                    PdfToolsService.remove_blank_pages(output_path, output_path)
                    return True, None

            # HTML -> DOCX
            elif target_ext == "docx":
                try:
                    from bs4 import BeautifulSoup
                    import docx
                    soup = BeautifulSoup(html_content, "html.parser")
                    doc = docx.Document()
                    title = soup.title.string if soup.title else "Converted HTML Document"
                    doc.add_heading(title, 0)

                    for elem in soup.find_all(["h1", "h2", "h3", "p", "li"]):
                        if elem.name == "h1":
                            doc.add_heading(elem.get_text(), level=1)
                        elif elem.name == "h2":
                            doc.add_heading(elem.get_text(), level=2)
                        elif elem.name == "h3":
                            doc.add_heading(elem.get_text(), level=3)
                        elif elem.name == "li":
                            doc.add_paragraph(elem.get_text(), style="List Bullet")
                        else:
                            doc.add_paragraph(elem.get_text())

                    doc.save(output_path)
                    if progress_callback:
                        progress_callback(100, "HTML converted to DOCX.")
                    return True, None
                except Exception as e:
                    return False, f"HTML to DOCX error: {str(e)}"

            else:
                return False, f"Unsupported HTML target format: {target_ext}"

        except Exception as e:
            return False, f"HTML conversion error: {str(e)}"
